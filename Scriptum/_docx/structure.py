#!/usr/bin/env python3
# coding: utf-8
#
# part of:
#   S C R I P T U M 
#
# by  temmel007@gmail.com
# 2020
#
###################
# MODULE _docx.structure
# PROVIDES 
#   a view on the document as is and filters for sections, paragraphs, tables, images
#   extracts templates
#   provides functions to find elements
#   even if it creates a high level structures view it delivers mostly
#   lists of [(path, tag, lowlevel elements), ...]
#   where path is a list like ['section:load_bc', 'subsection:load', 'head']
#         tag is a tag element
#         lowlevel element is a docx Paragraph or Table

from typing import Any, List, Tuple, Union, TYPE_CHECKING

from .template import copy_table_before, copy_paragraph_before, add_page_break_before

from ..tag.tag import Tag, getTag
from ..rdf.reportDataFile import docx_sections
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.table import CT_Tbl
from copy import deepcopy

if TYPE_CHECKING:
    from .images.element import DocImageBlockElement
    from .tables.element import DocTableBlockElement
    from .paragraphs.element import DocTextBlockElement

class StructuredElement:

    HEADER = 'STRUCTURE'

    def __init__(self, 
                 tag: Tag,    # the very first tag found that opens the structure
                 path: List,  # kind of address
                 parent: Any, # until we unterstand a Union here: Union[StructuredElement,], 
                 content: List[Tuple[Tag,Any]], # actually Any is an element for python-docx
                 anchor=None #Union[DocParagraphElement,None]
                 ):
        """
        BASE CLASS - not used directly

        everything is a structure
        - a section, a subsection, and so on
        - a table (addressable), otherwise it is a part of a structure
        - an image as structure or paragraph
        - a simple paragraph or a recoursive paragrapgh

        a structure has 
        - an opening paragraph
        - content
        - a closing paragraph

        the opening paragraph will be kept, while the tag is later deleted
        the closing paragraph (self.anchor.thing) will be kept during generation as you can always add new content before
        that paragraph, at the end it will be deleted (self.markForDeletion) - self.parent is required for this

        """

        # may contain errors or warnings
        self.error = []
        self.warning = []
        
        self.name = tag.name
        self.structure = []
        self.type = tag.ns
        self.path = path+[f'{tag.ns}:{tag.name}']

        # tags and active elements
        # main
        self.tag = tag
        # collect all tags in that structure
        self.tags = []
        self.active = []
        for _tag, elem in content:
            if type(_tag) == Tag:
                self.tags += [_tag]
                self.active += [elem]
            elif tag == 'struct':
                self.active += [elem]

        self.parent = parent
        self.anchor = anchor
        # for sections and subsections etc, we require more anchors
        self.subAnchors = []
        #print('build', parent, self.type, self.type in docx_sections['order'])
        if parent and self.type in docx_sections['order']: # not None
            self.parent.subAnchors += [content[0][1]]
            #print(len(self.parent.subAnchors))
        self.raw_content = content

        if path and anchor: # keep last paragraph in sections!
            anchor.anchor = True # mark that paragraph
            self.markForDeletion(anchor)

        # inside the structure we may have usable templates
        # is it marked as a template?
        self.isTemplate = 'template' in tag.args
        
        if len(path) > 0 and path[0] == 'section:template':
            # all is a template from here on
            self.isTemplate = True
        #print(tag.args,path,self.isTemplate)

    def explore(self):

        substructure = []
        subtag = None
        for etag, elem in self.raw_content:
            #print('explore a',etag,elem)
            #if etag:
            #    print(etag.puretag, etag.tagtype, tag.puretag,elem)
            if not etag:
                # simple element no tags or anything to change
                if not substructure:
                    # that element belongs to this structure
                    self.structure += [(etag,elem)]
                else:
                    # that element belongs to the next level substructure
                    substructure += [(etag,elem)]
            elif etag.tagtype == 'simple':
                if not substructure:
                    # that element belongs to this structure
                    # differ only for images, tables are always structures, see below
                    if etag.ns == 'image':
                        from .images.element import DocImageBlockElement
                        _newStructure = DocImageBlockElement(etag, self.path, self, [(etag,elem)], None)
                        #_newStructure.explore()
                        _newStructure.structure = [(etag,elem)]
                        self.structure += [('image',_newStructure)]
                    elif self.type == 'section' and self.name == 'template':
                        elem.path = self.path+[etag.puretag]
                        elem.isTemplate = True
                        self.structure += [(etag,elem)]
                    else:
                        # we just add this element anyhow
                        elem.path = self.path
                        self.structure += [(etag,elem)]
                        # in case this is marked as anchor we can reset the structure self.anchor
                        # this can help while adding substructures before that point just to have
                        # content correctly placed at the end of that structure
                        if 'anchor' in etag.args:
                            self.anchor = elem
                else:
                    # that element belongs to the next level substructure
                    substructure += [(etag,elem)]
            elif etag.puretag == self.tag.puretag:
                # keep this element as main entry and exit elements from and to that structure
                self.structure += [(etag,elem)]
                if etag.ns in ['table']:
                    self.markForDeletion(elem)
                # if etag.ns not in ['table']:
                #     # the opening or closing of that structure is part of the structure
                #     self.structure += [(etag,elem)]
                # else:
                #     print(etag.puretag,tag.puretag,etag.ns,tag.ns,elem)
                #     self.error += ['you cannot open or close that structure in a table: %s: %s, %s'%(self.path, etag.rawtag, tag.rawtag)]
            elif etag.tagtype == 'open':
                # a new open tag
                # why does this belong to the structure itself?
                #self.structure += [(etag,elem)]
                #self.markForDeletion(DocParagraphElement(elem))
                self.markForDeletion(elem)
                if not substructure:
                    # store the last opened tag to avoid mixtures
                    subtag = etag

                # if etag.ns in ['table']:
                #     #self.markForDeletion(DocParagraphElement(elem))
                #     self.markForDeletion(elem)
                # collect a new substructure
                substructure += [(etag,elem)]
            elif etag.tagtype == 'close':
                if not subtag:
                    self.error += ['There is a closed tag, not opened or a wrong simple tag: %s: %s'%(self.path, etag.rawtag)]
                    continue
                if etag.puretag == subtag.puretag:

                    substructure += [(etag,elem)]

                    if etag.ns == 'table':
                        self.markForDeletion(elem)
                        from .tables.element import DocTableBlockElement
                        _newStructure = DocTableBlockElement(subtag, self.path, self, substructure, anchor=None)
                        _newStructure.explore()
                        self.structure += [('table',_newStructure)]
                    elif etag.ns == 'image':
                        self.markForDeletion(elem)
                        from .images.element import DocImageBlockElement
                        _newStructure = DocImageBlockElement(subtag, self.path, self, substructure, anchor=None)
                        _newStructure.explore()
                        self.structure += [('image',_newStructure)]
                    elif self.type == 'section' and self.name == 'template':
                        from .paragraphs.element import DocTextBlockElement
                        _newStructure = DocTextBlockElement(subtag, self.path, self, substructure, anchor=None)
                        _newStructure.explore()
                        self.structure += [('text',_newStructure)]
                    else:
                        _newStructure = StructuredElement(subtag, self.path, self, substructure, anchor=elem)
                        _newStructure.explore()
                        self.structure += [('struct',_newStructure)]

                    
                    #self.structure += [(etag,elem)]
                    substructure = []
                    subtag = None
                elif subtag:
                    substructure += [(etag,elem)]
                else:
                    self.structure += [(etag,elem)]
                    if etag.ns == 'marker':
                        #self.markForDeletion(DocParagraphElement(elem))
                        self.markForDeletion(elem)
        if subtag:
            # there is some unclosed tag...
            self.error += [f'There is an open, but never closed tag, probably a typo: {self.path} - {subtag.rawtag}']
        #print('explore z', self.structure)
        #self.generateTemplates()

    def iterOnStructures(self):
        newiter = []
        for t,e in self.structure:
            if t in ['struct', 'image', 'table']:
                newiter += [e]+e.iterOnStructures()
        #yield from newiter
        return newiter

    def findExact(self, path):
        """expect a list with entries, that start with path
        we will found either one or nothing """

        l = len(path)
        result = []

        # this is any sub level
        # in order to avoid search inside all, we first look for the parent
        #print('search exact...:', l, self.path, path)
        for t,e in self.structure:
            #print('exact what?',self.path,t,e.path,path)
            # cycle over all elements in that structure
            if t in ['text','struct', 'table', 'image'] and e.path == path:
                # the thing I am looking for is a table or a image in that structure
                #print('exact t or i',t,e.path,path)
                result += [(t,e)]
                break
            elif type(t) == Tag and (e.path+[t.puretag])[:l] == path:
                # the thing I am looking for is ????
                #print('exact tag',(e.path+[t.puretag])[:l],'<->',path,e.path,t.puretag)
                result += [(t,e)]
                break
            elif e.path[:l] == path:
                # at least the thing and the element e have the same root? 
                #print('exact subpath',e.path[:l],'<->',path)
                result += [(t,e)]
                break
            # elif e.path == path:
            #     # maybe we find and mean that very structure itself?
            #     result += [(t,e)]
            #     break
            else:
                # nothing matched
                if 0:
                    if type(t) == Tag:
                        print(f"exact nothing {t.puretag} {e.path} <?> {path}")
                    else:
                        print(f'exact nothing {t} {e.path} <?> {path}')
                pass

        if not result:
            print(f'WARNING: No exact match: {path}')

        return result
    
    def __iter__(self):
        """iterate content to all levels, but keep table and image structures as is
        this is required for findGlobal only"""
        newiter = []
        for t,e in self.structure:
            #print('create iterator',t,e, e.path)
            if t in ['struct']:#, 'image', 'table']:
                for sub in e:
                    #print(sub)
                    newiter += [sub]
            else:
                newiter += [(t,e)]
        yield from newiter

    def findGlobal(self, path) -> list:
        found = []
        for t,e in self: # requires __iter__ above
            if not t or e.anchor: continue # never change anchors or simple paragraphs
            if t == 'text':
                continue # for now we ignore textblock inner tags
            elif t == 'table':
                for st, se in e.structure:
                    if not st or st.burned: continue
                    if st.puretag == path:
                        found += [(st,se)]
            elif t == 'image':
                if e.path[-1] == path:
                    found += [(t,e)]
            elif t.puretag == path:
                found += [(t,e)]
        return found

    def getTemplates(self):
        templates = []
        for t,e in self.structure:
            if t in ['struct', 'image', 'table']:
                if e.isTemplate:
                    e.createTemplate()
                    templates += [(t,e)]
                # always look into the rest as well
                templates += e.getTemplates()
            elif e.isTemplate:
                e.createTemplate()
                templates += [(t,e)]
        return templates
        
    def markForDeletion(self,elem):
        # store that on the highest level
        self.parent.markForDeletion(elem)

    def createTemplate(self):
        self.deepcopy = []
        #print('createTemplate',self.structure)
        done = []
        for tag, elem in self.structure:
            if elem in done: continue
            done.append(elem)
            #print(tag,elem)
            if type(tag) == str: # -> struct?
                self.deepcopy += elem.createTemplate()
            elif hasattr(elem.thing,'_tbl'):
                self.deepcopy += [deepcopy(elem.thing._tbl)]
            else:
                self.deepcopy += [deepcopy(elem.thing._p)]
        return self.deepcopy

    def copy(self, anchor, parent, newpath=[], newname='', section=None):
        """redfined in the subclasses, however here used to copy the structure allowever

        copy all elements just before the anchor,
        in case of self.type == 'struct', rename the first and last element if newname is set"""
        
        from .paragraphs import DocParagraphElement
        from .tables import DocTableElement

        newElements = []

        #print('anchor:',anchor.thing.text)

        # if we create a new structure and the first element posesses the 'breakbefore' tag
        # we add a new page_break
        try:
            if (text:= self.deepcopy[0].text):
                tag = getTag(text)[0]
                if 'breakbefore' in tag.args:
                    #print('add a page break...')
                    add_page_break_before(anchor.thing)
        except Exception as e:
            print(f'INFO: No page break due to {e}')

        for dc in self.deepcopy:
            if type(dc) == CT_Tbl:
                newElements += [
                    DocTableElement(
                        copy_table_before(anchor.thing,  
                                            deepcopy(dc)) # copy once again to keep template clean
                                        )
                                ]

            else:
                newElements += [
                    DocParagraphElement(
                        copy_paragraph_before(anchor.thing,  
                                            deepcopy(dc)) # copy once again to keep template clean
                                        )
                                ]
        
        if newname:

            obj = newElements[0] # first element is always a paragraph
            tag = obj.tags[0] # always first tag!
            obj.replaceTag(tag,f'<{newname}>')
            tag.rewriteTag(newname)
            
            obj = newElements[-1] # last element is always a paragraph
            tag = obj.tags[-1] # always last tag!
            obj.replaceTag(tag,f'</{newname}>')
            tag.rewriteTag(newname)

        # and finally unfold it again
        newUnfoldedElements = []
        for e in newElements:
            if e.tags:
                for subt in e.tags:
                    newUnfoldedElements += [(subt, e)]
            else:
                newUnfoldedElements += [(None, e)]

        # the last one is again the new anchor...
        # but skip subAnchor creations
        _subs = [s for s in parent.subAnchors]
        mycopy = StructuredElement(newElements[0].tags[0], newpath, parent, newUnfoldedElements, e)
        parent.subAnchors = _subs
        mycopy.explore()
        #print('mycopy', mycopy.path, 'into', parent.path, 'at', anchor.path)
        # add this one into the parents structure
        parent.structure.append(('struct',mycopy))
        #print(parent.path)
        #print(parent.structure)
        # and to the addressbook:
        # in this case we may have a series of new entries structures of structures with images and tables
        section.buildAddressbook()
        #print('MMMM',mycopy.path)
        #print(section.addressbook.keys())
        #section.addressbook['.'.join(mycopy.path)] = mycopy
        return mycopy

    def clean(self):
        print('clean() TO BE IMPLEMENTED IN SUBCLASSES')
        pass

    def delete(self):
        """delete the full structure
        
        access to deletion for final user
        """
        print(f'Deleting structure {(".".join(self.path))}')
        from .paragraphs import delete_paragraph
        from .tables import delete_table
        for t, e in self.structure:
            #print(el)
            if t in ['struct', 'image', 'table']:
                e.delete()
            elif type(e) == Paragraph:
                delete_paragraph(e)
            elif type(e) == Table:
                delete_table(e)
        print('... deleted')

    def inspect(self, level=None, indent=0):
        """visualize the content of that element"""
        _indent = ' '*indent
        print(f'{_indent}{self.HEADER} (start): {self.type}:{self.name}')
        print(f'{_indent}  path: {self.path}')
        print(f'{_indent}  parent: {self.parent}')
        print(f'{_indent}  me: {self}')
        
        if self.error:
            print(f'{_indent}  has errors:\n    '+'\n    '.join(self.error))
        print(f'{_indent}  will iterate as:')
        for t,e in self.structure:
            if t in ['struct', 'image', 'table']:
                if level == None:
                    e.inspect(indent=indent+3)
                elif level == 0:
                    print(f'{_indent}  skip next structure: {e.HEADER} {e.name}')
                else:
                    e.inspect(level=level-1, indent=indent+3)
            else:
                if type(t) == Tag:
                    print(f'{_indent}  tagged element: {t.rawtag} {e}')
                else:
                    if hasattr(e,'text'):
                        print(f'{_indent}  simple element: {t} {e} {e.text}')
                    else:
                        print(f'{_indent}  simple element: {t} {e}')
        if self.anchor:
            print(f'{_indent}  anchor: {self.anchor} {self.anchor.thing}')
        print(f'{_indent}{self.HEADER} (end): {self.name}')


