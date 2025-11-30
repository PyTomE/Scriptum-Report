"""Table-related element implementations for Word rendering."""

from __future__ import annotations

from copy import deepcopy
from typing import Any, List, Optional, Tuple

from docx.oxml.table import CT_Tbl

from ...element import TableElement
from ...tag import Tag, getTag
from ..base import DocElement
from ..paragraphs import DocParagraphElement
from ..structure import StructuredElement
from ..template import copy_paragraph_before, copy_table_before

class DocTableElement(DocElement, TableElement):
    """tables only 
    - tables with tags inside: variant container - default
    - tables as a whole: variant pure
    """
    def __init__(self, elem):
        # init 
        super().__init__(elem)
        self.type = 'table'
        self.content ='_table_'
        self.thing = elem
        self.path = []
        self.anchor = False

        tags = []
        for row in elem.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    obj = DocParagraphElement(p)
                    found = []
                    for tag in obj.tags:
                        if tag.tagtype != 'simple':
                            tag.tagtype = (
                                f'invalid:notallowed_{tag.tagtype}_tag_{tag.puretag}'
                            )  # allow only simple tags
                            found += [tag]
                        elif tag.ns == 'comment':
                            # ignore - better remove
                            obj.replaceTag(tag,'')
                            tag.burn()
                        else:
                            found += [tag]
                    tags += found
        self.tags = tags

    #
    # replace tags
    #
    def replaceTag(self, tag: Tag, replace: str) -> Optional[Any]:
        table = self.thing
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    obj = DocParagraphElement(paragraph)
                    #print('RT',paragraph.text,tag, obj)
                    run = obj.replaceTag(tag,replace)
                    if run is not None:
                        return run
        return None

def delete_table(table):
    """tested for docx only"""
    t = table._element
    if t.getparent() is not None:
        t.getparent().remove(t)
        t._tbl = t._element = None


class DocTableBlockElement(StructuredElement):
    """Table block representation, backed by surrounding paragraphs."""

    HEADER = "TABLE"

    def __init__(
        self,
        tag: Tag,
        path: List,
        parent: Any,
        content: List[Tuple[Tag, Any]],
        anchor: None,
    ) -> None:
        super().__init__(tag, path, parent, content, None)
        self.subtype = "structure"

    def copy(self, anchor, parent, newpath=None, newname="", section=None):
        """Copy the table block just before the anchor."""

        if newpath is None:
            newpath = []

        newElements = []
        for dc in self.deepcopy:
            if isinstance(dc, CT_Tbl):
                newElements += [
                    DocTableElement(
                        copy_table_before(anchor.thing, deepcopy(dc))
                    )
                ]

            else:
                newElements += [
                    DocParagraphElement(
                        copy_paragraph_before(anchor.thing, deepcopy(dc))
                    )
                ]

        if newname:
            obj = newElements[0]
            tag = obj.tags[0]
            obj.replaceTag(tag, f"<{newname}>")
            tag.rewriteTag(newname)

            obj = newElements[-1]
            tag = obj.tags[-1]
            obj.replaceTag(tag, f"</{newname}>")
            tag.rewriteTag(newname)

        newUnfoldedElements = []
        for element in newElements:
            if element.tags:
                for subt in element.tags:
                    newUnfoldedElements += [(subt, element)]
            else:
                newUnfoldedElements += [(None, element)]

        mycopy = DocTableBlockElement(
            newElements[0].tags[0], newpath, parent, newUnfoldedElements, None
        )
        mycopy.explore()
        parent.structure.append(("table", mycopy))
        section.addressbook[".".join(mycopy.path)] = mycopy
        return mycopy

    def fill(self, task) -> None:
        """Fill the element with a value."""

        value = task.value
        actions = task.actions
        value.load()

        if value.object.exists:
            for t, table in self.structure:
                if isinstance(table, DocTableElement):
                    break

            value.object.fill(table)

            for key, val in actions.items():
                #print('\nKV',key,val)
                for t, element in self.structure:
                    #print('te',t,element)
                    if not t or t.burned:
                        continue
                    #print('pure',key, t.puretag)
                    if key == t.puretag:
                        element.replaceTag(t, str(val))
                        t.burn()
        else:
            leading_par = self.structure[0][1]
            leading_par.replaceTag(self.tag, str(value.object))

    def clean(self) -> None:
        """Clean the unused tags and paragraphs."""

        for t, element in self.structure:
            if not t or t.burned:
                continue
            element.replaceTag(t, "")
            t.burn()
            element.deleteIfEmpty()

